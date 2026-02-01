"""
Фаза 4, Неделя 4: мультимодальность — изображения и документы.

Обработка изображений: Moondream Station (MLX) → Ollama vision (moondream/llava).
Обработка документов: PDF, DOCX → извлечение текста для RAG.
Рекомендации: Backend (единая точка маршрутизации), SRE (таймауты, fallback).
"""

import base64
import io
import logging
import os
from pathlib import Path
from typing import Optional, Tuple

import httpx

from app.config import get_settings

logger = logging.getLogger(__name__)

# Изображения
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"}
IMAGE_CONTENT_TYPES = {"image/png", "image/jpeg", "image/gif", "image/webp", "image/bmp"}

# Документы (10+ форматов: PDF, DOCX, DOC, TXT, HTML, ODT, RTF и др.)
DOCUMENT_EXTENSIONS = {
    ".pdf", ".docx", ".doc",
    ".txt", ".text",
    ".html", ".htm", ".xhtml",
    ".odt", ".ods", ".odp",
    ".rtf",
}
DOCUMENT_CONTENT_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
    "text/plain", "text/html", "application/xhtml+xml",
    "application/vnd.oasis.opendocument.text",
    "application/rtf", "text/rtf",
}


def _get_moondream_url() -> str:
    return getattr(get_settings(), "moondream_station_url", None) or os.getenv("MOONDREAM_STATION_URL", "http://localhost:2020")


def _get_ollama_url() -> str:
    return get_settings().ollama_url


def detect_content_type(filename: Optional[str] = None, content_type: Optional[str] = None) -> str:
    """
    Определение типа контента для маршрутизации.
    Возвращает: "image" | "document" | "unknown".
    """
    if content_type:
        ct = content_type.lower().split(";")[0].strip()
        if ct in IMAGE_CONTENT_TYPES:
            return "image"
        if ct in DOCUMENT_CONTENT_TYPES:
            return "document"
    if filename:
        ext = Path(filename).suffix.lower()
        if ext in IMAGE_EXTENSIONS:
            return "image"
        if ext in DOCUMENT_EXTENSIONS:
            return "document"
    return "unknown"


class MultimodalProcessor:
    """
    Единая точка обработки изображений и документов для RAG и чата.
    Изображения: Moondream Station → Ollama vision. Документы: извлечение текста.
    """

    def __init__(
        self,
        moondream_url: Optional[str] = None,
        ollama_url: Optional[str] = None,
        vision_enabled: bool = True,
        vision_timeout: float = 60.0,
    ):
        self.moondream_url = (moondream_url or _get_moondream_url()).rstrip("/")
        self.ollama_url = (ollama_url or _get_ollama_url()).rstrip("/")
        self.vision_enabled = vision_enabled
        self.vision_timeout = vision_timeout

    async def process_image(
        self,
        image_base64: Optional[str] = None,
        image_path: Optional[str] = None,
        prompt: str = "Опиши это изображение подробно.",
    ) -> Optional[str]:
        """
        Обработка изображения через Vision: Moondream Station → Ollama.
        image_base64: base64-строка (или data:image/...;base64,...).
        image_path: путь к файлу (альтернатива base64).
        """
        if not self.vision_enabled:
            logger.debug("Vision disabled, skipping image processing")
            return None

        raw_b64: Optional[str] = None
        if image_base64:
            raw_b64 = image_base64.split(",")[-1].strip() if "," in image_base64 else image_base64
        elif image_path and Path(image_path).exists():
            try:
                with open(image_path, "rb") as f:
                    raw_b64 = base64.b64encode(f.read()).decode("utf-8")
            except Exception as e:
                logger.warning("Failed to read image file: %s", e)
                return None
        if not raw_b64:
            logger.warning("No image data provided")
            return None

        # 1) Moondream Station REST API
        try:
            async with httpx.AsyncClient(timeout=self.vision_timeout) as client:
                r = await client.post(
                    f"{self.moondream_url}/v1/query",
                    json={"image": raw_b64, "prompt": prompt},
                )
                if r.status_code == 200:
                    data = r.json()
                    answer = data.get("answer", "").strip()
                    if answer:
                        logger.info("Vision: Moondream Station OK")
                        return answer
        except Exception as e:
            logger.debug("Moondream Station failed: %s", e)

        # 2) Ollama vision fallback
        for model in ("moondream", "llava:7b", "moondream:latest"):
            try:
                async with httpx.AsyncClient(timeout=self.vision_timeout) as client:
                    r = await client.post(
                        f"{self.ollama_url}/api/generate",
                        json={
                            "model": model,
                            "prompt": prompt,
                            "images": [raw_b64],
                            "stream": False,
                        },
                    )
                    if r.status_code == 200:
                        data = r.json()
                        text = data.get("response", "").strip()
                        if text:
                            logger.info("Vision: Ollama %s OK", model)
                            return text
            except Exception as e:
                logger.debug("Ollama %s failed: %s", model, e)
        logger.warning("All vision backends failed")
        return None

    def extract_document_text(
        self,
        file_path: Optional[str] = None,
        content: Optional[bytes] = None,
        filename_hint: Optional[str] = None,
    ) -> Optional[str]:
        """
        Извлечение текста из документов для RAG (10+ форматов).
        Поддерживаются: PDF, DOCX, DOC, TXT, HTML, ODT, RTF.
        Опциональные зависимости: PyMuPDF, python-docx.
        """
        ext = ""
        if file_path and Path(file_path).exists():
            ext = Path(file_path).suffix.lower()
            if ext == ".pdf":
                return self._extract_pdf(file_path=file_path)
            if ext in (".docx", ".doc"):
                return self._extract_docx(file_path=file_path)
            if ext in (".txt", ".text"):
                return self._extract_txt(file_path=file_path)
            if ext in (".html", ".htm", ".xhtml"):
                return self._extract_html(file_path=file_path)
            if ext == ".odt":
                return self._extract_odt(file_path=file_path)
            if ext == ".rtf":
                return self._extract_rtf(file_path=file_path)
        if content:
            ext = (Path(filename_hint or "").suffix.lower()) if filename_hint else ""
            if ext == ".pdf" or content[:4] == b"%PDF":
                return self._extract_pdf(content=content)
            # PK = zip: DOCX имеет [Content_Types].xml, ODT — content.xml
            if content[:2] == b"PK":
                try:
                    import zipfile
                    z = zipfile.ZipFile(io.BytesIO(content), "r")
                    names = z.namelist()
                    z.close()
                    if "content.xml" in names and ext != ".docx":
                        return self._extract_odt(content=content)
                    if "[Content_Types].xml" in names or ext in (".docx", ".doc"):
                        return self._extract_docx(content=content)
                except Exception:
                    pass
            elif ext in (".docx", ".doc") or b"[Content_Types].xml" in content[:65536]:
                return self._extract_docx(content=content)
            if ext in (".txt", ".text") or (not ext and b"<" not in content[:200] and content[:2] != b"PK"):
                return self._extract_txt(content=content)
            if ext in (".html", ".htm", ".xhtml") or content[:20].strip().lower().startswith(b"<!doctype") or content[:10].strip().lower().startswith(b"<html"):
                return self._extract_html(content=content)
            if ext == ".odt":
                return self._extract_odt(content=content)
            if ext == ".rtf" or content[:5] == b"{\\rtf":
                return self._extract_rtf(content=content)
        return None

    def _extract_pdf(self, file_path: Optional[str] = None, content: Optional[bytes] = None) -> Optional[str]:
        try:
            import fitz  # PyMuPDF
        except ImportError:
            logger.debug("PyMuPDF not installed; pip install pymupdf for PDF extraction")
            return None
        try:
            if content:
                doc = fitz.open(stream=content, filetype="pdf")
            else:
                doc = fitz.open(file_path)
            parts = []
            for page in doc:
                parts.append(page.get_text())
            doc.close()
            return "\n\n".join(p.strip() for p in parts if p.strip())
        except Exception as e:
            logger.warning("PDF extraction failed: %s", e)
            return None

    def _extract_docx(self, file_path: Optional[str] = None, content: Optional[bytes] = None) -> Optional[str]:
        try:
            from docx import Document
        except ImportError:
            logger.debug("python-docx not installed; pip install python-docx for DOCX extraction")
            return None
        try:
            if content:
                doc = Document(io.BytesIO(content))
            else:
                doc = Document(file_path)
            return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            logger.warning("DOCX extraction failed: %s", e)
            return None

    def _extract_txt(self, file_path: Optional[str] = None, content: Optional[bytes] = None) -> Optional[str]:
        """TXT: чтение текста (built-in, без зависимостей)."""
        try:
            if content:
                return content.decode("utf-8", errors="replace").strip()
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                return f.read().strip()
        except Exception as e:
            logger.warning("TXT extraction failed: %s", e)
            return None

    def _extract_html(self, file_path: Optional[str] = None, content: Optional[bytes] = None) -> Optional[str]:
        """HTML: извлечение текста (built-in html.parser)."""
        import html.parser

        class TextExtractor(html.parser.HTMLParser):
            def __init__(self):
                super().__init__()
                self.text = []

            def handle_data(self, data):
                s = data.strip()
                if s:
                    self.text.append(s)

        try:
            raw = content.decode("utf-8", errors="replace") if content else Path(file_path).read_text(encoding="utf-8", errors="replace")
            parser = TextExtractor()
            parser.feed(raw)
            return "\n\n".join(parser.text) if parser.text else None
        except Exception as e:
            logger.warning("HTML extraction failed: %s", e)
            return None

    def _extract_odt(self, file_path: Optional[str] = None, content: Optional[bytes] = None) -> Optional[str]:
        """ODT: OpenDocument Text (zip + content.xml)."""
        import zipfile
        import xml.etree.ElementTree as ET

        try:
            if content:
                z = zipfile.ZipFile(io.BytesIO(content), "r")
            else:
                z = zipfile.ZipFile(file_path, "r")
            raw = z.read("content.xml").decode("utf-8", errors="replace")
            z.close()
            root = ET.fromstring(raw)
            ns = {"text": "urn:oasis:names:tc:opendocument:xmlns:text:1.0"}
            parts = [el.text or "" for el in root.iter() if el.text and el.text.strip()]
            parts += [el.tail or "" for el in root.iter() if el.tail and el.tail.strip()]
            return "\n\n".join(p.strip() for p in parts if p.strip())
        except Exception as e:
            logger.warning("ODT extraction failed: %s", e)
            return None

    def _extract_rtf(self, file_path: Optional[str] = None, content: Optional[bytes] = None) -> Optional[str]:
        """RTF: базовая декодировка (удаление тегов)."""
        import re
        try:
            raw = content.decode("utf-8", errors="replace") if content else Path(file_path).read_text(encoding="utf-8", errors="replace")
            text = re.sub(r"\\[a-z]+\d*\s?", " ", raw)
            text = re.sub(r"[{}]", " ", text)
            text = re.sub(r"\s+", " ", text).strip()
            return text if text else None
        except Exception as e:
            logger.warning("RTF extraction failed: %s", e)
            return None

    async def get_text_for_rag(
        self,
        content_type: str,
        image_base64: Optional[str] = None,
        image_path: Optional[str] = None,
        document_path: Optional[str] = None,
        document_content: Optional[bytes] = None,
        prompt_for_image: str = "Опиши содержимое для поиска по базе знаний. Выдели ключевые факты и термины.",
    ) -> Optional[str]:
        """
        Единая точка для пайплайна: по типу контента возвращает текст для RAG.
        """
        if content_type == "image":
            return await self.process_image(
                image_base64=image_base64,
                image_path=image_path,
                prompt=prompt_for_image,
            )
        if content_type == "document":
            if document_path:
                return self.extract_document_text(file_path=document_path)
            if document_content:
                return self.extract_document_text(content=document_content)
        return None


_multimodal_processor: Optional[MultimodalProcessor] = None


def get_multimodal_processor() -> MultimodalProcessor:
    """Singleton для использования в роутерах и RAG."""
    global _multimodal_processor
    if _multimodal_processor is None:
        s = get_settings()
        _multimodal_processor = MultimodalProcessor(
            moondream_url=getattr(s, "moondream_station_url", None),
            ollama_url=s.ollama_url,
            vision_enabled=getattr(s, "multimodal_vision_enabled", True),
            vision_timeout=getattr(s, "multimodal_vision_timeout", 60.0),
        )
    return _multimodal_processor
